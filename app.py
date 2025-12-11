import streamlit as st
import pdfplumber
from docx import Document  # from python-docx
import os
from dotenv import load_dotenv
from openai import OpenAI
import io
from fpdf import FPDF
import math
import stripe
from urllib.parse import urlparse, parse_qs

import json
import firebase_admin
from firebase_admin import credentials, auth as fb_auth, firestore


# Configure the Streamlit page
st.set_page_config(page_title="DocSift", layout="wide")


# Simple upgrade URL placeholder – change this later to your real Stripe checkout or pricing page
UPGRADE_URL = "https://buy.stripe.com/14A7sLej62Qw7Xs2vsbZe00"

# Dev flag so you (and only you) can enable a Pro toggle in dev
DEV_MODE = os.getenv("DOCSIFT_DEV_MODE", "false").lower() == "true"


# Stripe API key from secrets
stripe.api_key = st.secrets["STRIPE_SECRET_KEY"]

# ---------- Firebase Admin initialization (for plan verification) ----------

@st.cache_resource
def init_firebase():
    """
    Initialize Firebase Admin SDK once and reuse.
    """
    if not firebase_admin._apps:
        # Load service account JSON from Streamlit secrets
        service_account_info = json.loads(st.secrets["firebase"]["service_account_json"])
        cred = credentials.Certificate(service_account_info)
        firebase_admin.initialize_app(cred)
    return firestore.client()

db = init_firebase()


def sync_plan_from_firebase():
    """
    If a Firebase ID token is present in the URL (?token=...),
    verify it, look up the user doc in Firestore, and set
    st.session_state.is_pro_user accordingly.

    We key Firestore docs by lowercased email, matching the landing page.
    """
    # Handle Streamlit query params for both new and older APIs
    try:
        params = st.query_params  # Streamlit 1.40+
    except AttributeError:
        params = st.experimental_get_query_params()

    raw_token = params.get("token", None)
    if raw_token is None:
        # No token → nothing to do; allow anonymous usage
        return

    # raw_token may be a list or a string depending on API
    if isinstance(raw_token, list):
        id_token = raw_token[0]
    else:
        id_token = raw_token

    if not id_token:
        return

    try:
        decoded = fb_auth.verify_id_token(id_token)
    except Exception:
        # If token is invalid/expired, treat as anonymous user
        return

    email = (decoded.get("email") or "").lower().strip()
    if not email:
        return

    # Store email in session for display
    st.session_state["firebase_email"] = email

    # Look up Firestore user doc: users/{emailKey}
    doc_ref = db.collection("users").document(email)
    snap = doc_ref.get()

    plan = "free"
    if snap.exists:
        data = snap.to_dict()
        plan = (data.get("plan") or "free").lower()

    st.session_state["firebase_plan"] = plan

    # Synchronize Pro/Free into the app's flag
    st.session_state.is_pro_user = (plan == "pro")


# Run once at import time so the plan is ready before UI renders
sync_plan_from_firebase()




# -------- Usage Tracking --------
if "daily_uses" not in st.session_state:
    st.session_state.daily_uses = 0

if "is_pro_user" not in st.session_state:
    st.session_state.is_pro_user = False  # default: free user

if "usage_log" not in st.session_state:
    st.session_state.usage_log = []


# -------- Fixed header + sticky tabs CSS --------
st.markdown(
    """
    <style>
    /* Fixed header bar below the Streamlit toolbar */
    .docsift-fixed-header {
        position: fixed;
        top: 3.2rem;             /* below Streamlit's top toolbar */
        left: 18rem;             /* shift right so it's not under the sidebar */
        right: 0;
        z-index: 1000;
        padding: 0.6rem 3rem 0.7rem 3rem;
        background-color: #0E1117;  /* Streamlit dark background */
        color: #FFFFFF;
        box-shadow: 0 2px 4px rgba(0,0,0,0.4);
    }

    /* Optional: when the sidebar is collapsed on smaller screens,
       let the header stretch full width again */
    @media (max-width: 1000px) {
        .docsift-fixed-header {
            left: 0;
        }
        
    }

    .docsift-fixed-header h1 {
        margin: 0;
        font-size: 2.2rem;   /* larger, real title */
        font-weight: 700;
        color: #FFFFFF;
    }

    .docsift-fixed-header p {
        margin: 0.2rem 0 0 0;
        font-size: 0.95rem;
        opacity: 0.85;
        color: #FFFFFF;
    }

    
    /* Push main content down so it isn't hidden behind header + tabs */
    .block-container {
        padding-top: 11.0rem;   /* toolbar + header + tabs */
    }

    /* Optional: widen main content area */
    .main .block-container {
        max-width: 2000px !important;
        padding-left: 3rem;
        padding-right: 3rem;
    }
    </style>
    """,
    unsafe_allow_html=True,
)



# Render fixed header once (outside sidebar / main logic)
st.markdown(
    """
    <div class="docsift-fixed-header">
        <h1>DocSift</h1>
        <p>Your AI-powered document analysis toolkit.</p>
    </div>
    """,
    unsafe_allow_html=True,
)





# --------- Safety limits (we'll later make these per-plan) ---------
FREE_MAX_PAGES = 5
PRO_MAX_PAGES = 50   # or 100 if you want

APPROX_CHARS_PER_PAGE = 1800    # For DOCX/TXT page estimates
# -------------------------------------------------------------------


# Load environment variables from .env
load_dotenv()


# Read the API key from the environment
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")


# Create an OpenAI client instance
client = OpenAI(api_key=OPENAI_API_KEY)


def check_limits(page_count):
    # Pro users get a higher limit
    if st.session_state.is_pro_user:
        if page_count > PRO_MAX_PAGES:
            return f"Pro accounts can process documents up to {PRO_MAX_PAGES} pages."
        return None

    # Free user limits
    if page_count > FREE_MAX_PAGES:
        return f"Free accounts can only process documents up to {FREE_MAX_PAGES} pages."

    if st.session_state.daily_uses >= 3:
        return "Free accounts can only analyze 3 documents per day."

    return None


def render_pro_upsell(feature_name: str):
    """Standard Pro upsell block shown when a feature is locked."""
    st.info(
        f"🔒 **{feature_name} is a DocSift Pro feature.**\n\n"
        "Upgrade to DocSift Pro to unlock this and all advanced analysis tools "
        "(Key Points, Action Items, Risk Analyzer, Explain Like I'm 12, "
        "Clarity Rewrite, Study Guides, and Full Reports)."
    )
    st.markdown(f"[👉 Upgrade to DocSift Pro]({UPGRADE_URL})")


def estimate_pages_from_text(text: str) -> int:
    """Roughly estimate page count from text length for DOCX/TXT files.
    This is a simple approximation just to enforce safety limits.
    """
    if not text:
        return 1
    pages = math.ceil(len(text) / APPROX_CHARS_PER_PAGE)
    return max(1, pages)

def chunk_text(text: str, max_chars: int = 4000, overlap: int = 500) -> list[str]:
    """Split a long string into overlapping chunks of up to max_chars.
    Overlap helps keep context between chunks.
    """
    chunks = []
    start = 0
    text_length = len(text)

    while start < text_length:
        end = start + max_chars
        chunk = text[start:end]

        chunks.append(chunk)

        # Move start forward with overlap
        start = end - overlap
        if start < 0:
            start = 0

    return chunks



def extract_text_from_pdf(file) -> str:
    """Extract text from a PDF file (limits enforced later via check_limits)."""
    with pdfplumber.open(file) as pdf:
        
        text = ""
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text += page_text + "\n\n"

    return text.strip()


def extract_text_from_docx(file) -> str:
    """Extract text from a DOCX file (limits enforced later via check_limits)."""
    doc = Document(file)
    full_text = "\n".join([para.text for para in doc.paragraphs])

    return full_text


def extract_text_from_txt(file) -> str:
    """Extract raw text from a TXT file (limits enforced later via check_limits)."""
    
    # Read raw bytes from the uploaded file
    raw = file.read()
    
    # Attempt UTF-8 decode first, fall back to latin-1
    try:
        text = raw.decode("utf-8", errors="ignore")
    except AttributeError:
        # In case it's already a Python string (rare)
        text = raw
    
    text = text.strip()

    return text
    

def summarize_text(text: str) -> str:
    """Use the OpenAI API to summarize the given text."""
    # Optionally, truncate the text so we don't send extremely long inputs
    max_chars = 6000
    if len(text) > max_chars:
        text_to_summarize = text[:max_chars]
    else:
        text_to_summarize = text

    # Create a prompt/instruction for the model
    system_message = (
        "You are an assistant that summarizes documents. "
        "Produce a clear, concise summary in a few short paragraphs. "
        "Focus on the main ideas and key points only."
    )


    user_message = (
        "Here is the document text. Summarize it:\n\n"
        f"{text_to_summarize}"
    )


    try:
        response = client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=[
                {"role": "system", "content": system_message},
                {"role": "user", "content": user_message},
            ],
            temperature=0.3,
        )
        # Extract the assistant's reply text
        summary = response.choices[0].message.content
        return summary.strip()
    
    except Exception as e:
        # If something goes wrong, return an error message we can display
        return f"[Error summarizing text: {e}]"


def summarize_document(text: str) -> str:
    """Summarize the document, using chunking for longer texts.
    This keeps costs predictable while handling longer docs better.
    """

    # If short enough, just use the existing summarizer
    max_chars_single = 4000
    if len(text) <= max_chars_single:
        return summarize_text(text)

    # For longer docs, break into chunks
    chunks = chunk_text(text, max_chars=max_chars_single, overlap=500)

    # Safety cap: don't process too many chunks in one go
    max_chunks = 5
    if len(chunks) > max_chunks:
        chunks = chunks[:max_chunks]

    partial_summaries = []
    for idx, chunk in enumerate(chunks, start=1):
        # Summarize each chunk using the existing summarizer
        chunk_summary = summarize_text(chunk)
        partial_summaries.append(f"Chunk {idx} summary:\n{chunk_summary}")

    # Combine chunk summaries and summarize again to get a final overall summary
    combined = "\n\n".join(partial_summaries)
    final_summary = summarize_text(combined)

    return final_summary.strip()


def extract_key_points(text: str) -> str:
    """Extract key bullet-point ideas from the document using AI."""

    system_message = (
        "You extract key ideas from a document. "
        "Return the top 8–12 most important points as clear, concise bullet points. "
        "Do not summarize with paragraphs — ONLY bullet points."
    )

    # Truncate long text to avoid high token usage for now
    max_chars = 6000
    if len(text) > max_chars:
        text_to_process = text[:max_chars]
    else:
        text_to_process = text

    user_message = (
        "Extract the key points from the following document:\n\n"
        f"{text_to_process}"
    )

    try:
        response = client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=[
                {"role": "system", "content": system_message},
                {"role": "user", "content": user_message},
            ],
            temperature=0.2,
        )
        result = response.choices[0].message.content
        return result.strip()

    except Exception as e:
        return f"[Error extracting key points: {e}]"


def extract_action_items(text: str) -> str:
    """Extract action items, tasks, and follow-ups from the document using AI."""

    system_message = (
        "You analyze documents and extract action items. "
        "Return a list of specific tasks, decisions, and follow-ups. "
        "For each item, include who is responsible if mentioned, and any deadlines or dates. "
        "Use bullet points. Be concise and practical."
    )

    # Limit size for now to keep requests cheap and fast
    max_chars = 6000
    if len(text) > max_chars:
        text_to_process = text[:max_chars]
    else:
        text_to_process = text

    user_message = (
        "Extract all action items, tasks, decisions, and follow-ups from the following document.\n\n"
        "If no explicit owner is given, just describe the task.\n\n"
        f"{text_to_process}"
    )

    try:
        response = client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=[
                {"role": "system", "content": system_message},
                {"role": "user", "content": user_message},
            ],
            temperature=0.2,
        )
        result = response.choices[0].message.content
        return result.strip()

    except Exception as e:
        return f"[Error extracting action items: {e}]"


def extract_risks(text: str) -> str:
    """Identify risks, issues, and red flags in the document using AI."""

    system_message = (
        "You analyze documents for potential risks and red flags. "
        "Highlight anything that could be a problem, confusing, incomplete, or high-risk. "
        "Group them into categories if helpful (e.g., Legal, Financial, Operational, Clarity). "
        "Use bullet points with short explanations."
    )

    max_chars = 6000
    if len(text) > max_chars:
        text_to_process = text[:max_chars]
    else:
        text_to_process = text

    user_message = (
        "Read the following document and identify any potential risks, issues, or red flags. "
        "This may include unclear responsibilities, missing data, legal concerns, unrealistic assumptions, "
        "or anything that might cause problems later.\n\n"
        f"{text_to_process}"
    )

    try:
        response = client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=[
                {"role": "system", "content": system_message},
                {"role": "user", "content": user_message},
            ],
            temperature=0.2,
        )
        result = response.choices[0].message.content
        return result.strip()

    except Exception as e:
        return f"[Error extracting risks: {e}]"


def explain_like_12(text: str) -> str:
    """Rewrite the document in simple language that a 12-year-old could understand."""

    system_message = (
        "You explain complex documents in simple, friendly language. "
        "Your job is to rewrite the content so that a typical 12-year-old could understand it. "
        "Use short sentences, everyday words, and clear examples. "
        "Avoid jargon. If you must use a technical term, briefly explain what it means."
    )

    max_chars = 6000
    if len(text) > max_chars:
        text_to_process = text[:max_chars]
    else:
        text_to_process = text

    user_message = (
        "Rewrite the following text so that a 12-year-old could understand it. "
        "Keep the main ideas, but make it much simpler and clearer:\n\n"
        f"{text_to_process}"
    )

    try:
        response = client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=[
                {"role": "system", "content": system_message},
                {"role": "user", "content": user_message},
            ],
            temperature=0.4,
        )
        result = response.choices[0].message.content
        return result.strip()

    except Exception as e:
        return f"[Error simplifying text: {e}]"


def rewrite_for_clarity(text: str) -> str:
    """Rewrite the document to be clearer, more concise, and professional."""

    system_message = (
        "You are an expert editor who improves clarity, conciseness, and professionalism. "
        "Rewrite the text so it is easy to read, well-organized, and uses a neutral, professional tone. "
        "Do not remove important information. Do not change the meaning. "
        "Avoid jargon unless necessary, and break up very long sentences."
    )

    # Limit how much we send in one go to keep costs predictable
    max_chars = 6000
    if len(text) > max_chars:
        text_to_process = text[:max_chars]
    else:
        text_to_process = text

    user_message = (
        "Rewrite the following text to be clearer, more concise, and more professional, "
        "while preserving the original meaning:\n\n"
        f"{text_to_process}"
    )

    try:
        response = client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=[
                {"role": "system", "content": system_message},
                {"role": "user", "content": user_message},
            ],
            temperature=0.3,
        )
        result = response.choices[0].message.content
        return result.strip()

    except Exception as e:
        return f"[Error rewriting for clarity: {e}]"


def generate_study_guide(text: str) -> str:
    """Create a study guide and quiz-style questions from the document using AI."""

    system_message = (
        "You turn documents into study guides. "
        "Create a clear, structured study guide that helps someone learn this material. "
        "Use markdown formatting with headings and bullet points. "
        "Include key concepts, questions WITH answers, and a brief summary."
    )

    max_chars = 6000
    if len(text) > max_chars:
        text_to_process = text[:max_chars]
    else:
        text_to_process = text

    user_message = (
        "Create a study guide from the following document. "
        "Structure it with these sections:\n"
        "1. Key Concepts (bullet points)\n"
        "2. Questions and Answers (at least 8–12 Q&A pairs)\n"
        "3. Short Summary (one short paragraph)\n\n"
        "Here is the document:\n\n"
        f"{text_to_process}"
    )

    try:
        response = client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=[
                {"role": "system", "content": system_message},
                {"role": "user", "content": user_message},
            ],
            temperature=0.3,
        )
        result = response.choices[0].message.content
        return result.strip()

    except Exception as e:
        return f"[Error generating study guide: {e}]"


def build_full_report(text: str) -> str:
    """Build a full markdown report using all of DocSift's analysis helpers."""

    # Call each of the existing helper functions
    summary = summarize_document(text)
    key_points = extract_key_points(text)
    action_items = extract_action_items(text)
    risks = extract_risks(text)
    eli12 = explain_like_12(text)
    study_guide = generate_study_guide(text)

    # Stitch everything together as a single markdown string
    sections = [
        "# DocSift Report",
        "Generated by DocSift.\n",
        "## 1. Summary",
        summary,
        "## 2. Key Points",
        key_points,
        "## 3. Action Items",
        action_items,
        "## 4. Risks / Red Flags",
        risks,
        "## 5. Explain Like I'm 12",
        eli12,
        "## 6. Study Guide",
        study_guide,
    ]

    # Join all the pieces with blank lines in between
    full_report_markdown = "\n\n".join(sections)
    return full_report_markdown



def generate_pdf_from_markdown(markdown_text: str) -> bytes:
    """Generate a simple PDF file (as bytes) from the given markdown text.
    This version treats the markdown as plain text and sanitizes unsupported characters.
    """

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    pdf.set_font("Helvetica", size=12)

    for line in markdown_text.splitlines():
        # Safely convert unsupported characters
        safe_line = line.encode("latin-1", "replace").decode("latin-1")
        pdf.multi_cell(0, 8, safe_line)
        pdf.ln(1)

    # Get PDF output (could be str OR bytearray depending on fpdf version)
    pdf_bytes = pdf.output(dest="S")

    # If it's a string, encode it
    if isinstance(pdf_bytes, str):
        pdf_bytes = pdf_bytes.encode("latin-1")

    # If it's a bytearray, convert to bytes
    if isinstance(pdf_bytes, bytearray):
        pdf_bytes = bytes(pdf_bytes)

    return pdf_bytes


def generate_docx_from_markdown(markdown_text: str) -> bytes:
    """Generate a Word .docx file (as bytes) from the given markdown text."""

    doc = Document()

    for line in markdown_text.splitlines():
        line = line.rstrip()

        if not line:
            # Blank line → blank paragraph
            doc.add_paragraph("")
            continue

        # Headings based on leading #
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        # Bullet points starting with - or *
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            # Regular paragraph
            doc.add_paragraph(line)

    # Save to an in-memory bytes buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()




# ---- Streamlit app layout ----
# ------------- Sidebar: upload & info -------------
with st.sidebar:

    st.header("Upload a Document")

    uploaded_file = st.file_uploader(
        "Choose a PDF, Word (.docx), or Text (.txt) file",
        type=["pdf", "docx", "txt"]
    )

    st.markdown("**Current safety limits**")
    st.caption(
        f"- Free: up to **{FREE_MAX_PAGES}** pages per document\n"
        f"- Pro: up to **{PRO_MAX_PAGES}** pages per document\n"
        f"- Very large files and books are not supported."
    )

    st.markdown("---")
    st.subheader("Account")

    # NEW: show Firebase-linked account info if available
    firebase_email = st.session_state.get("firebase_email")
    firebase_plan = st.session_state.get("firebase_plan")

    if firebase_email:
        st.caption(f"Signed in as: {firebase_email}")

    if firebase_plan:
        st.caption(f"Account plan: {firebase_plan.capitalize()}")


    # ---- Pro User ----
    if st.session_state.is_pro_user:
        st.markdown("### ⭐ DocSift Pro Active")
        st.caption("Thank you for supporting DocSift!")
    
    # ---- Free User ----
    else:
        st.markdown(
            """
            <div style="line-height:1.5;">
                <a href="https://getdocsift.com/#account" target="_blank"
                   style="font-weight:600; text-decoration:none;">
                   Upgrade to DocSift Pro
                </a>
                <br>
                Unlock advanced features including:<br>
                <em>Key Points, Action Items, Risk Analysis, Explain Like I'm 12,<br>
                Clarity Rewrites, Study Guides, and Full Reports.</em>
            </div>
            """,
            unsafe_allow_html=True,
        )
        
    # --- Dev-only override (does NOT reset is_pro_user in production) ---

    if DEV_MODE:
        st.markmarkdown("---")
        st.caption("Developer controls (visible only in DEV_MODE).")
        dev_force_pro = st.checkbox(
            "Force Pro mode (dev only)",
            value=st.session_state.get("is_pro_user", False),
        )
        if dev_force_pro:
            st.session_state.is_pro_user = True




# ------------- Main content -------------

if uploaded_file is None:
    # No file yet: show a simple welcome panel
    st.subheader("Welcome to DocSift")
    st.write(
        "Upload a document using the sidebar to begin. "
        "DocSift will extract the text and help you summarize, analyze risks, "
        "find action items, simplify language, and build a full report."
    )

else:
    # Try to extract text from the uploaded file, with safety limits
    text = ""
    error_message = None

    # Determine file extension (more reliable than MIME in some cases)
    file_name = uploaded_file.name.lower()
    if file_name.endswith(".pdf"):
        file_type = "pdf"
    elif file_name.endswith(".docx"):
        file_type = "docx"
    elif file_name.endswith(".txt"):
        file_type = "txt"
    else:
        file_type = "unknown"

    try:
        if file_type == "pdf":
            text = extract_text_from_pdf(uploaded_file)
        elif file_type == "docx":
            text = extract_text_from_docx(uploaded_file)
        elif file_type == "txt":
            text = extract_text_from_txt(uploaded_file)
        else:
            raise ValueError("Unsupported file type. Please upload PDF, DOCX, or TXT.")
    except Exception as e:
        error_message = str(e)

    if error_message:
        st.error(error_message)
    else:
        # At this point, text is ready and passed safety limits
        st.success("Document loaded successfully.")

        # ---------- STEP 5: Usage & Free-Tier Limits ----------
        # Estimate pages from text (for DOCX/TXT) or from PDF earlier logic
        page_count = estimate_pages_from_text(text)

        # Check free/pro limits
        limit_error = check_limits(page_count)
        if limit_error:
            st.error(limit_error)
            st.stop()

        # Count this use for free-tier users only
        if not st.session_state.is_pro_user:
            st.session_state.daily_uses += 1

        # Create tabbed interface for all analysis tools
        tab_overview, tab_summary, tab_keypoints, tab_actions, tab_risks, tab_eli12, tab_rewrite, tab_study, tab_full = st.tabs([
            "Overview",
            "Summary",
            "Key Points",
            "Action Items",
            "Risks",
            "Explain Like I'm 12",
            "Rewrite for Clarity",
            "Study Guide",
            "Full Report"
        ])

        
        # ---- Overview tab: show raw extracted text ----
        with tab_overview:
            st.subheader("Extracted Text")
            st.write("This tab shows the raw text extracted from your document. Review it here before generating summaries or reports.")

            
            st.text_area("Document Text", value=text, height=400)

                # ---- Summary tab ----
        with tab_summary:
            st.subheader("AI Summary")
            st.write("Generate a clean, concise summary of your document. This provides a quick high-level understanding without reading the entire text.")

            if st.button("Generate Summary"):
                with st.spinner("Summarizing document..."):
                    summary = summarize_document(text)
                st.write(summary)

                st.session_state.usage_log.append({"action": "summary"})


                # ---- Key Points tab ----
        with tab_keypoints:
            st.subheader("Key Points")
            # STEP 9: Upgrade prompt
            if not st.session_state.is_pro_user:
                render_pro_upsell("Key Points")


            else:
                st.write(
                    "Extract the most important points, ideas, and takeaways from the document. "
                    "Useful for quick reviews or preparing presentations."
                )

                if st.button("Generate Key Points"):
                    with st.spinner("Finding key ideas..."):
                        key_points = extract_key_points(text)
                    st.write(key_points)

                    st.session_state.usage_log.append({"action": "key_points"})


                # ---- Action Items tab ----
        with tab_actions:
            st.subheader("Action Items")

            if not st.session_state.is_pro_user:
                render_pro_upsell("Action Items")

            else:
                st.write(
                    "Identify actionable steps, tasks, and follow-ups based on the document content. "
                    "Perfect for meetings, project planning, and workflows."
                )

                if st.button("Generate Action Items"):
                    with st.spinner("Looking for tasks and follow-ups..."):
                        action_items = extract_action_items(text)
                    st.write(action_items)

                    st.session_state.usage_log.append({"action": "action_items"})


                # ---- Risks / Red Flags tab ----
        with tab_risks:
            st.subheader("Risks / Red Flags")

            if not st.session_state.is_pro_user:
                render_pro_upsell("Risk Analyzer")

            
            else:
                st.write(
                    "Automatically scan the document for risks, gaps, red flags, or potential issues. "
                    "Ideal for contracts, policies, audits, and compliance reviews."
                )

                if st.button("Analyze Risks"):
                    with st.spinner("Scanning for risks and red flags..."):
                        risks = extract_risks(text)
                    st.write(risks)

                    st.session_state.usage_log.append({"action": "risks"})


                # ---- Explain Like I'm 12 tab ----
        with tab_eli12:
            st.subheader("Explain Like I'm 12")

            if not st.session_state.is_pro_user:
                render_pro_upsell("Explain Like I'm 12")

            
            else:
                st.write(
                    "Simplify complex text into clear, easy-to-understand language. "
                    "Great for learning, teaching, or breaking down technical documents."
                )

                if st.button("Simplify the Document"):
                    with st.spinner("Rewriting in simple language..."):
                        simple_version = explain_like_12(text)
                    st.write(simple_version)

                    st.session_state.usage_log.append({"action": "eli12"})


                # ---- Rewrite for Clarity tab ----
        with tab_rewrite:
            st.subheader("Rewrite for Clarity & Professional Tone")

            if not st.session_state.is_pro_user:
                render_pro_upsell("Rewrite for Clarity")

            
            else:
                st.write(
                    "Improve clarity, professionalism, and readability while preserving meaning. "
                    "Ideal for emails, reports, policies, and long-form content."
                )

                if st.button("Rewrite Document for Clarity"):
                    with st.spinner("Rewriting document for clarity and professional tone..."):
                        rewritten = rewrite_for_clarity(text)
                    st.write(rewritten)

                    st.session_state.usage_log.append({"action": "rewrite_clarity"})


               # ---- Study Guide / Quiz tab ----
        with tab_study:
            st.subheader("Study Guide / Quiz")

            if not st.session_state.is_pro_user:
                render_pro_upsell("Study Guide / Quiz")

            
            else:
                st.write(
                    "Turn your document into a structured study guide with key points, lessons, and "
                    "quiz questions. Perfect for training, learning, and onboarding."
                )

                if st.button("Generate Study Guide"):
                    with st.spinner("Creating study guide and quiz questions..."):
                        study_guide = generate_study_guide(text)
                    st.markdown(study_guide)

                    st.session_state.usage_log.append({"action": "study_guide"})


                # ---- Full Report tab ----
        with tab_full:
            st.subheader("Full Report (All Sections)")

            if not st.session_state.is_pro_user:
                render_pro_upsell("Full Report")

            
            else:
                st.write(
                    "Generate a complete analysis—including summary, key points, actions, risks, "
                    "simplified explanation, clarity rewrite, and study guide—all in one "
                    "comprehensive report."
                )

                if st.button("Generate Full Report"):
                    with st.spinner(
                        "Generating full report (summary, key points, actions, risks, ELI12, "
                        "clarity rewrite, study guide)..."
                    ):
                        st.session_state["full_report"] = build_full_report(text)

                        st.session_state.usage_log.append({"action": "full_report"})

                full_report = st.session_state.get("full_report")

                if full_report:
                    pdf_bytes = generate_pdf_from_markdown(full_report)
                    docx_bytes = generate_docx_from_markdown(full_report)

                    st.markdown("### Download Full Report")

                    st.download_button(
                        label="Download Full Report (Markdown)",
                        data=full_report,
                        file_name="docsift_report.md",
                        mime="text/markdown",
                        key="download_md_top",
                    )

                    st.download_button(
                        label="Download Full Report (PDF)",
                        data=pdf_bytes,
                        file_name="docsift_report.pdf",
                        mime="application/pdf",
                        key="download_pdf_top",
                    )

                    st.download_button(
                        label="Download Full Report (Word)",
                        data=docx_bytes,
                        file_name="docsift_report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="download_docx_top",
                    )

                    st.markdown("---")
                    st.markdown(full_report)
                    st.markdown("---")

                    st.markdown("### Download Full Report")

                    st.download_button(
                        label="Download Full Report (Markdown)",
                        data=full_report,
                        file_name="docsift_report.md",
                        mime="text/markdown",
                        key="download_md_bottom",
                    )

                    st.download_button(
                        label="Download Full Report (PDF)",
                        data=pdf_bytes,
                        file_name="docsift_report.pdf",
                        mime="application/pdf",
                        key="download_pdf_bottom",
                    )

                    st.download_button(
                        label="Download Full Report (Word)",
                        data=docx_bytes,
                        file_name="docsift_report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="download_docx_bottom",
                    )

